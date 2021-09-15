import argparse

from utils import load
from docx import Document


def is_journal(pub):
    return 'journal' in pub['bib'] and 'rxiv' not in pub['bib']['journal'].lower()

def is_preprint(pub):
    return 'journal' in pub['bib'] and 'rxiv' in pub['bib']['journal'].lower()

def is_conference(pub):
    return 'conference' in pub['bib']

def is_other(pub):
    return not is_journal(pub) and not is_preprint(pub) and not is_conference(pub)


def docx(author, pubs):

    # split publications in sections
    journals = filter(is_journal, pubs)
    preprints = filter(is_preprint, pubs)
    conferences = filter(is_conference, pubs)
    other = filter(is_other, pubs)

    # sort by date
    sort_params = dict(key=lambda x: x['bib'].get('pub_year', 0), reverse=True)
    journals = sorted(journals, **sort_params)
    preprints = sorted(preprints, **sort_params)
    conferences = sorted(conferences, **sort_params)
    other = sorted(other, **sort_params)

    import pdb; pdb.set_trace()

    # helper function
    def _publication_entry(pub, document):
        p = document.add_paragraph()
        data = pub['bib']

        # title line
        title = data['title']
        title = p.add_run(title)
        title.bold = True
        title.add_break()

        # authors line
        authors = data['author']
        authors = authors.split(' and ')

        def format_author(full_name):  # Fabio Carrara
            names = full_name.split(' ')
            given_names = names[:-1]
            last_name = names[-1]

            abbr = [f'{n[0]}.' for n in given_names]
            abbr = ' '.join(abbr)

            return f'{last_name}, {abbr}'  # Carrara, F.

        authors = list(map(format_author, authors))
        authors = '; '.join(authors)
        authors = p.add_run(authors)
        authors.add_break()

        # venue / journal line
        pub_year = str(data.get('pub_year', ''))
        pub_year = p.add_run(pub_year)
        pub_year.underline = True
        p.add_run(', ')

        venue = data.get('journal', None) or data.get('conference', None)
        if venue:
            venue = p.add_run(venue)
            venue.italic = True

        if 'volume' in data:
            volume = data['volume']
            p.add_run(', ')
            volume = p.add_run(f'n. {volume}')
        
        if 'number' in data:
            number = data['number']
            number = p.add_run(f'({number})')

        if 'pages' in data:
            pages = data['pages']
            p.add_run(', ')
            pages = p.add_run(f'pp. {pages}')
        
        if 'publisher' in data:
            publisher = data['publisher']
            p.add_run(', ')
            publisher = p.add_run(publisher)

    # build the docx
    document = Document()
    document.add_heading(f"{author['name']}'s Publications", level=0)

    ## journals
    document.add_heading('Journal Papers', level=1)
    for pub in journals:
        _publication_entry(pub, document)

    ## conferences
    document.add_heading('Conference Papers', level=1)
    for pub in conferences:
        _publication_entry(pub, document)
    
    ## preprints
    document.add_heading('Preprints', level=1)
    for pub in preprints:
        _publication_entry(pub, document)

    ## others
    document.add_heading('Others', level=1)
    for pub in other:
        _publication_entry(pub, document)

    return document


def main(args):
    gs_author = load('google_scholar_author.pkl')
    gs_pubs = load('google_scholar_publications.pkl')

    if args.format == 'docx':
        document = docx(gs_author, gs_pubs)
        output = args.output or 'publication_list.docx'
        document.save(output)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Generate Publication List')
    parser.add_argument('format', default='docx', choices=('docx',), help='Output format')
    parser.add_argument('-o', '--output', type=str, help='Output path')

    args = parser.parse_args()
    main(args)